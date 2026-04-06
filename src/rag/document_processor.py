"""
文档处理器。

Description:
    解析各种格式的文档，提取文本内容用于知识库构建。
@author ganjianfei
@version 1.0.0
2026-04-05
"""

import json
import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from src.rag.chunker import SemanticChunker

logger = logging.getLogger(__name__)


@dataclass
class ParsedDocument:
    """解析后的文档。

    Attributes:
        title: 文档标题。
        content: 文档内容。
        metadata: 文档元数据。
        source: 来源文件名。
        doc_type: 文档类型。
        category: 商品类目。
    """

    title: str
    content: str
    metadata: dict[str, Any]
    source: str
    doc_type: str = "general"
    category: str | None = None


class DocumentProcessor:
    """文档处理器。

    解析多种格式的文档并分块。

    支持格式:
    - Markdown (.md)
    - Text (.txt)
    - JSON (.json)
    - PDF (.pdf) - 需要 pypdf
    - DOCX (.docx) - 需要 python-docx

    Example:
        >>> processor = DocumentProcessor()
        >>> doc = processor.parse("brand_guide.md", doc_type="brand_guide")
        >>> chunks = processor.process(doc)
    """

    def __init__(self) -> None:
        """初始化文档处理器。"""
        self.chunker = SemanticChunker()

    def parse(
        self,
        file_path: str | Path,
        doc_type: str = "general",
        category: str | None = None,
        metadata: dict[str, Any] | None = None,
    ) -> ParsedDocument:
        """解析文档文件。

        Args:
            file_path: 文件路径。
            doc_type: 文档类型。
            category: 商品类目。
            metadata: 额外元数据。

        Returns:
            解析后的文档。

        Raises:
            ValueError: 不支持的文件格式。
        """
        path = Path(file_path)
        suffix = path.suffix.lower()

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # 根据文件类型选择解析方法
        if suffix == ".md":
            content, title = self._parse_markdown(path)
        elif suffix == ".txt":
            content, title = self._parse_text(path)
        elif suffix == ".json":
            content, title = self._parse_json(path)
        elif suffix == ".pdf":
            content, title = self._parse_pdf(path)
        elif suffix == ".docx":
            content, title = self._parse_docx(path)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")

        # 合并元数据
        full_metadata = metadata or {}
        full_metadata["file_size"] = path.stat().st_size
        full_metadata["file_suffix"] = suffix

        logger.info(f"Parsed document: {path.name} ({len(content)} chars)")

        return ParsedDocument(
            title=title or path.stem,
            content=content,
            metadata=full_metadata,
            source=path.name,
            doc_type=doc_type,
            category=category,
        )

    def parse_content(
        self,
        content: str,
        title: str,
        doc_type: str = "general",
        category: str | None = None,
        metadata: dict[str, Any] | None = None,
    ) -> ParsedDocument:
        """直接解析文本内容。

        Args:
            content: 文档内容。
            title: 文档标题。
            doc_type: 文档类型。
            category: 商品类目。
            metadata: 额外元数据。

        Returns:
            解析后的文档。
        """
        return ParsedDocument(
            title=title,
            content=content,
            metadata=metadata or {},
            source="direct_input",
            doc_type=doc_type,
            category=category,
        )

    def process(self, document: ParsedDocument) -> list[dict[str, Any]]:
        """处理文档并生成分块。

        Args:
            document: 解析后的文档。

        Returns:
            分块字典列表，每个包含 content 和 metadata。
        """
        chunks = self.chunker.split_by_semantic(
            document.content,
            metadata={
                "title": document.title,
                "doc_type": document.doc_type,
                "category": document.category,
                "source": document.source,
                **document.metadata,
            },
        )

        return [
            {
                "content": chunk.content,
                "metadata": chunk.metadata,
            }
            for chunk in chunks
        ]

    def _parse_markdown(self, path: Path) -> tuple[str, str]:
        """解析 Markdown 文件。

        Args:
            path: 文件路径。

        Returns:
            (内容, 标题) 元组。
        """
        content = path.read_text(encoding="utf-8")

        # 尝试从第一个标题提取标题
        title = path.stem
        for line in content.split("\n"):
            if line.startswith("# "):
                title = line[2:].strip()
                break

        return content, title

    def _parse_text(self, path: Path) -> tuple[str, str]:
        """解析纯文本文件。

        Args:
            path: 文件路径。

        Returns:
            (内容, 标题) 元组。
        """
        content = path.read_text(encoding="utf-8")
        title = path.stem

        # 尝试从第一行提取标题
        first_line = content.split("\n")[0].strip()
        if first_line and len(first_line) < 100:
            title = first_line

        return content, title

    def _parse_json(self, path: Path) -> tuple[str, str]:
        """解析 JSON 文件。

        Args:
            path: 文件路径。

        Returns:
            (内容, 标题) 元组。
        """
        data = json.loads(path.read_text(encoding="utf-8"))

        # 提取标题
        title = data.get("title", data.get("name", path.stem))

        # 将 JSON 转换为文本
        if isinstance(data, dict):
            # 尝试提取主要内容字段
            main_fields = ["content", "text", "body", "description"]
            for field in main_fields:
                if field in data:
                    content = data[field]
                    if isinstance(content, str):
                        return content, title

            # 否则将整个 JSON 格式化为文本
            content = json.dumps(data, ensure_ascii=False, indent=2)
        elif isinstance(data, list):
            # 处理列表类型
            content = "\n\n".join(
                json.dumps(item, ensure_ascii=False, indent=2)
                if isinstance(item, dict)
                else str(item)
                for item in data
            )
        else:
            content = str(data)

        return content, title

    def _parse_pdf(self, path: Path) -> tuple[str, str]:
        """解析 PDF 文件。

        Args:
            path: 文件路径。

        Returns:
            (内容, 标题) 元组。

        Raises:
            ImportError: pypdf 未安装。
        """
        try:
            from pypdf import PdfReader
        except ImportError as e:
            raise ImportError("pypdf not installed. Run: pip install pypdf") from e

        reader = PdfReader(path)
        content = "\n\n".join(page.extract_text() or "" for page in reader.pages)

        # 尝试从元数据提取标题
        title = path.stem
        if reader.metadata and reader.metadata.title:
            title = reader.metadata.title

        return content.strip(), title

    def _parse_docx(self, path: Path) -> tuple[str, str]:
        """解析 DOCX 文件。

        Args:
            path: 文件路径。

        Returns:
            (内容, 标题) 元组。

        Raises:
            ImportError: python-docx 未安装。
        """
        try:
            from docx import Document
        except ImportError as e:
            raise ImportError("python-docx not installed. Run: pip install python-docx") from e

        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = "\n\n".join(paragraphs)

        # 尝试从第一个段落提取标题
        title = path.stem
        if paragraphs:
            first_para = paragraphs[0].strip()
            if len(first_para) < 100:
                title = first_para

        return content, title
